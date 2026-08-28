import os
import json
from docx import Document
from docxtpl import DocxTemplate
from jinja2 import Environment, FileSystemLoader
from playwright.sync_api import sync_playwright
from openai import OpenAI
from dotenv import load_dotenv

# Load environment variables (like OPENAI_API_KEY) from .env file
load_dotenv()

def read_docx(file_path):
    """Reads text and tables from a .docx file."""
    doc = Document(file_path)
    full_text = []
    
    # Extract regular paragraphs
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
            
    # Extract tables
    for table in doc.tables:
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                text = cell.text.strip().replace('\n', ' ')
                if text:
                    row_data.append(text)
            if row_data:
                full_text.append(" | ".join(row_data))
                
    return '\n'.join(full_text)

def generate_summary(text):
    """Uses OpenAI API to summarize the text based on specific headers."""
    client = OpenAI() # Assumes OPENAI_API_KEY is in environment

    system_prompt = """
You are an expert payroll processor and technical writer. 
Your task is to read the provided SOP document and extract/summarize information dynamically into a JSON object.

Follow these instructions carefully to populate the JSON output:

1. `title`: Extract the "Process Name" from the document and use it as the title (e.g., "Backdated PTO Process"). If not found, use a descriptive title based on the document content.
2. `process_overview`: Create a nested JSON object with keys "System", "Trigger", "TAT", "Accuracy".
   - For "System", look for "Software / Application Used" and extract it.
   - For "Trigger", strictly search the document for the exact phrase "Triggered By" (even if inside a table) and extract the text immediately following or associated with it (e.g. "Emails").
   - Extract "TAT" and "Accuracy" if they exist; otherwise set to null.
3. `dynamic_sections`: Locate the "INDEX" section in the document, specifically looking under the "Procedure with Screen Shots" part. Extract the list of sub-procedures (e.g., "Check Backdated PTO Requests", "Create Special Payroll", etc.). For each of these sub-procedures, read the corresponding content in the document and provide a concise summary or short bullet points. The output must be an array of objects, each containing a `header` (the sub-procedure name) and `summary` (string or array of bullet points).
4. `quick_checklist`: Generate a high-level array of checklist items by summarizing the content found specifically in the "Procedure Overview" section of the document.
5. `special_cases`: A short summary string or array of bullet points for special cases if they exist. Use key: `special_cases`. If none, set to null.
6. `full_sop_link`: Extract if present, else null. Use key: `full_sop_link`.
7. `frequency`: Extract if present, else null. Use key: `frequency`.
8. `contact`: Extract if present, else null. Use key: `contact`.

Ensure the output is strictly in JSON format. Do not copy and paste exact text for summaries; synthesize the information. Note that if information like TAT, Accuracy, Frequency, or Contact is missing from the DOCX source text, you must output null for them.
"""

    response = client.chat.completions.create(
        model="gpt-4o", # You can change this to gpt-3.5-turbo if needed
        messages=[
            {"role": "system", "content": system_prompt},
            {"role": "user", "content": f"Here is the SOP document text:\n\n{text}"}
        ],
        response_format={ "type": "json_object" },
        temperature=0.2
    )

    return json.loads(response.choices[0].message.content)

def save_as_json(data, output_path):
    """Saves a dictionary to a JSON file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        json.dump(data, f, indent=4)
    print(f"Saved JSON summary to {output_path}")

def save_as_txt(data, output_path):
    """Saves a dictionary to a formatted TXT file."""
    with open(output_path, 'w', encoding='utf-8') as f:
        for key, value in data.items():
            f.write(f"{key}\n")
            f.write("-" * len(key) + "\n")
            if isinstance(value, list):
                for item in value:
                    f.write(f"- {item}\n")
            elif isinstance(value, dict):
                for sub_key, sub_val in value.items():
                    f.write(f"  * {sub_key}: {sub_val}\n")
            else:
                f.write(f"{value}\n")
            f.write("\n")
    print(f"Saved TXT summary to {output_path}")

def save_as_docx(data, template_path, output_path):
    """Uses docxtpl to fill a docx template with dictionary data."""
    doc = DocxTemplate(template_path)
    
    # Render the template with the provided context dictionary
    doc.render(data)
    
    # Save the populated docx to the output path
    doc.save(output_path)
    print(f"Saved DOCX filled template to {output_path}")

def save_as_html(data, template_path, output_path):
    """Uses jinja2 to fill an HTML template with dictionary data."""
    template_dir, template_name = os.path.split(template_path)
    env = Environment(loader=FileSystemLoader(template_dir))
    template = env.get_template(template_name)
    
    html_out = template.render(data)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        f.write(html_out)
    print(f"Saved HTML filled template to {output_path}")

def save_as_pdf(html_path, pdf_path):
    """Uses playwright to convert an HTML file to PDF."""
    # Convert absolute or relative HTML path to a file:// URI
    file_uri = f"file:///{os.path.abspath(html_path).replace(chr(92), '/')}"
    
    with sync_playwright() as p:
        # Launch headless chromium
        browser = p.chromium.launch()
        page = browser.new_page()
        # Navigate to the local HTML file
        page.goto(file_uri)
        # Wait for fonts/styles to load (if any network resources are used)
        page.wait_for_load_state("networkidle")
        
        # Save as PDF (print background to keep colors)
        page.pdf(path=pdf_path, format="A4", print_background=True)
        browser.close()
    print(f"Saved PDF to {pdf_path}")

def merge_docx_files(master_path, append_path, output_path):
    """Merges two docx files, appending the second to the first with a page break."""
    from docxcompose.composer import Composer
    from docx import Document
    
    master = Document(master_path)
    master.add_page_break()
    
    composer = Composer(master)
    doc2 = Document(append_path)
    
    composer.append(doc2)
    composer.save(output_path)
    print(f"Merged {master_path} and {append_path} into {output_path}")

if __name__ == "__main__":
    input_docx = r"c:\Users\Intern\summary & chatbot\SOP_Emplova_Backdated PTO_V1.0.docx"
    
    # Create output directory based on the input filename
    base_name = os.path.splitext(os.path.basename(input_docx))[0]
    out_dir = os.path.join(r"c:\Users\Intern\summary & chatbot\output", base_name)
    os.makedirs(out_dir, exist_ok=True)

    extracted_text_path = os.path.join(out_dir, "extracted_text.txt")
    output_json = os.path.join(out_dir, "summary.json")
    output_txt = os.path.join(out_dir, "summary.txt")

    print("Reading DOCX file...")
    try:
        document_text = read_docx(input_docx)
        
        # Save the extracted text to a local file for inspection
        with open(extracted_text_path, "w", encoding="utf-8") as f:
            f.write(document_text)
        print(f"Saved exact text read by LLM to {extracted_text_path}")
        
    except Exception as e:
        print(f"Error reading docx: {e}")
        exit(1)

    print("Generating summary via OpenAI...")
    try:
        summary_data = generate_summary(document_text)
    except Exception as e:
        print(f"Error generating summary: {e}")
        exit(1)

    print("Saving outputs...")
    save_as_json(summary_data, output_json)
    save_as_txt(summary_data, output_txt)
    
    # Use standard templates for standalone run
    template_docx = r"c:\Users\Intern\summary & chatbot\input\template_grid.docx"
    output_docx = os.path.join(out_dir, "summary.docx")
    if os.path.exists(template_docx):
        save_as_docx(summary_data, template_docx, output_docx)
    else:
        print(f"Template not found at {template_docx}, skipping save_as_docx")
        
    template_html = r"c:\Users\Intern\summary & chatbot\input\template.html"
    output_html = os.path.join(out_dir, "summary.html")
    output_pdf = os.path.join(out_dir, "summary.pdf")
    
    if os.path.exists(template_html):
        save_as_html(summary_data, template_html, output_html)
        # Also generate PDF
        save_as_pdf(output_html, output_pdf)
    else:
        print(f"HTML Template not found at {template_html}, skipping save_as_html and save_as_pdf")
        
    print("Done!")
